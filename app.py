import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
from docx.shared import Pt
import time
import os

# -------------------------------------------------
# Page + theme
# -------------------------------------------------
st.set_page_config(page_title="AI Research Lab Simulator", page_icon="🧠", layout="centered")
st.markdown("""
    <style>
        body { background: radial-gradient(circle at top left, #000000, #0d0d0d 70%); color: white; }
        .stApp { background: rgba(10,10,10,0.9); backdrop-filter: blur(18px); }
        .block-container { padding-top: 3rem; max-width: 900px; }
        .title-text { text-align:center; font-size:2.6em; font-weight:800; color:#fff; text-shadow:0 0 25px rgba(0,255,255,0.4); }
        .sub-text { text-align:center; color:#bbb; margin-bottom: 1.2rem; }
        .stTextInput label { color:#fff !important; font-weight:600; }
        .stTextInput input { background-color:rgba(30,30,30,0.92); color:#fff; border-radius:10px; padding:0.6rem; border:1px solid rgba(0,255,255,0.25); }
        .stTextInput input::placeholder { color:#999; }
        .stButton button {
            background:linear-gradient(90deg,#00c6ff,#0072ff); color:white; font-weight:700; border-radius:12px;
            padding:0.6rem 1.3rem; border:none; box-shadow:0 0 20px rgba(0,114,255,0.4); transition:all 0.25s;
        }
        .stButton button:hover { transform:scale(1.04); box-shadow:0 0 30px rgba(0,198,255,0.6); }
        .agent-box {
            background:rgba(25,25,25,0.9); border:1px solid rgba(0,255,255,0.22);
            border-radius:15px; padding:1.0rem; margin-bottom:14px; color:#f2f2f2;
            box-shadow:0 0 15px rgba(0,255,255,0.08);
        }
        .stProgress > div > div > div > div { background-color:#00c6ff !important; }
    </style>
""", unsafe_allow_html=True)

st.markdown("<p class='title-text'>🧠 AI Research Lab Simulator</p>", unsafe_allow_html=True)
st.markdown("<p class='sub-text'>Gemini 2.5 Pro Multi-Agent Workflow • Chat • DOCX Export</p>", unsafe_allow_html=True)

# -------------------------------------------------
# State init
# -------------------------------------------------
if "paper" not in st.session_state:      # final, integrated paper text (for chat + export)
    st.session_state.paper = ""
if "summary" not in st.session_state:
    st.session_state.summary = ""
if "all_sections" not in st.session_state:  # concatenated outputs of agents
    st.session_state.all_sections = ""
if "chat_messages" not in st.session_state:  # chat history
    st.session_state.chat_messages = [
        {"role": "assistant", "content": "Hi! I’m your AI Research Assistant. Ask anything about the generated paper."}
    ]
if "model_ready" not in st.session_state:
    st.session_state.model_ready = False

# -------------------------------------------------
# Inputs
# -------------------------------------------------
api_key = st.text_input("🔑 Enter Your Gemini API Key", type="password", placeholder="Paste your Gemini API key")
topic = st.text_input("📘 Enter Your Research Topic", placeholder="e.g., Impact of AI on Precision Agriculture")

colA, colB = st.columns([1, 1])
with colA:
    do_generate = st.button("🚀 Generate Full Research Paper")
with colB:
    do_export = st.button("💾 Download .DOCX (Final Paper)")

# -------------------------------------------------
# Configure model (only when key present)
# -------------------------------------------------
model = None
if api_key:
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-pro")
        st.session_state.model_ready = True
    except Exception as e:
        st.warning(f"API init error: {e}")

# -------------------------------------------------
# Agent prompts (functions return text)
# -------------------------------------------------
def run_researcher(topic):
    prompt = (
        f"You are a professional scientific researcher. Generate a detailed, structured research paper draft on '{topic}'. "
        "Include: Abstract, Introduction, Literature Review, Methodology, Results/Discussion, and Conclusion. "
        "Be rigorous, well-organized, and use academic tone."
    )
    return model.generate_content(prompt).text

def run_reviewer(draft):
    prompt = (
        "You are a senior peer reviewer. Review the following draft for logic, structure, and factual soundness. "
        "List key feedback, then provide an improved version with revisions applied.\n\n"
        f"{draft}"
    )
    return model.generate_content(prompt).text

def run_editor(reviewed):
    prompt = (
        "You are an academic editor. Polish grammar, clarity, and transitions; ensure a consistent scholarly tone. "
        "Return only the improved paper text.\n\n"
        f"{reviewed}"
    )
    return model.generate_content(prompt).text

def run_citation_builder(edited):
    prompt = (
        "You are a citation builder. Insert realistic in-text citations (Author, Year) where appropriate and append a 'References' section. "
        "Avoid fabricating specific paper titles/DOIs. Keep them realistic but not falsified.\n\n"
        f"{edited}"
    )
    return model.generate_content(prompt).text

def run_fact_verifier(cited):
    prompt = (
        "You are a fact-verification agent. Check key factual statements and gently correct or clarify anything that seems inaccurate, "
        "keeping the paper cohesive. Return the verified, final text.\n\n"
        f"{cited}"
    )
    return model.generate_content(prompt).text

def run_summarizer(verified):
    prompt = (
        "Create a concise, professional 200-word executive summary of the following paper:\n\n"
        f"{verified}"
    )
    return model.generate_content(prompt).text

def run_research_composer(all_sections, topic):
    prompt = (
        f"You are a Research Composer integrating all segments into a single publication-ready paper on '{topic}'. "
        "Produce a unified document with sections: Title Page, Abstract, Introduction, Literature Review, Methodology, "
        "Results and Discussion, Conclusion, References. Ensure flow, consistency, and remove redundancies.\n\n"
        f"All prior agent outputs:\n{all_sections}"
    )
    return model.generate_content(prompt).text

# -------------------------------------------------
# DOCX export utility
# -------------------------------------------------
def make_docx(title, summary, final_text) -> bytes:
    doc = Document()
    # Title
    doc.add_heading(title or "AI Research Lab Simulator — Final Paper", 0)
    # Summary
    if summary:
        doc.add_heading("Executive Summary", level=1)
        p = doc.add_paragraph(summary)
        p.runs[0].font.size = Pt(11)
    # Final paper
    doc.add_heading("Research Paper", level=1)
    for line in (final_text or "").splitlines():
        doc.add_paragraph(line)
    # Save to memory
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# -------------------------------------------------
# GENERATION PIPELINE
# -------------------------------------------------
if do_generate:
    if not api_key or not topic:
        st.warning("Please enter both Gemini API Key and research topic.")
    elif not st.session_state.model_ready:
        st.warning("Model is not ready. Check your API key.")
    else:
        st.session_state.paper = ""
        st.session_state.summary = ""
        st.session_state.all_sections = ""

        stages = [
            ("🧠 Researcher", run_researcher),
            ("🔍 Reviewer", run_reviewer),
            ("✍️ Editor", run_editor),
            ("🧩 Citation Builder", run_citation_builder),
            ("🧪 Fact Verifier", run_fact_verifier),
            ("🧭 Summarizer", run_summarizer),
            ("📚 Research Composer", None),  # will run at the end using all sections
        ]

        progress = st.progress(0.0)
        last_output = ""
        concat_sections = ""

        for idx, (name, func) in enumerate(stages):
            st.subheader(f"{name}")
            with st.spinner(f"{name} is working..."):
                if name == "📚 Research Composer":
                    result = run_research_composer(concat_sections, topic)
                    st.session_state.paper = result  # final unified paper
                elif name == "🧭 Summarizer":
                    result = func(last_output)
                    st.session_state.summary = result
                    concat_sections += f"\n\n--- {name} Output ---\n{result}"
                    last_output = result
                else:
                    # normal stage
                    result = func(topic) if name == "🧠 Researcher" else func(last_output)
                    concat_sections += f"\n\n--- {name} Output ---\n{result}"
                    last_output = result

            st.markdown(f"<div class='agent-box'>{result}</div>", unsafe_allow_html=True)
            progress.progress((idx + 1) / len(stages))
            time.sleep(0.2)

        # Save sections for chat reference
        st.session_state.all_sections = concat_sections

        st.success("🎉 Complete Research Paper Generated!")
        st.info("Scroll down to chat with the AI about this paper, or export as .DOCX.")

# -------------------------------------------------
# DOCX DOWNLOAD
# -------------------------------------------------
if do_export:
    if not st.session_state.paper:
        st.warning("Generate the paper first, then export.")
    else:
        doc_bytes = make_docx(
            title=f"Research Paper — {topic}" if topic else "Research Paper",
            summary=st.session_state.summary,
            final_text=st.session_state.paper
        )
        st.download_button(
            label="⬇️ Download .DOCX",
            data=doc_bytes,
            file_name="Final_Research_Paper.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

st.divider()

# -------------------------------------------------
# CHAT (stable, no resets)
# -------------------------------------------------
st.header("💬 Chat with Your AI Research Assistant")
st.caption("Ask anything about your generated paper. The assistant will use the latest paper content for answers.")

if not st.session_state.model_ready:
    st.info("Enter your API key above to enable chat.")
else:
    # Render chat history
    for msg in st.session_state.chat_messages:
        with st.chat_message("assistant" if msg["role"] == "assistant" else "user"):
            st.markdown(msg["content"])

    # Chat input
    user_msg = st.chat_input("Type your question…")
    if user_msg:
        st.session_state.chat_messages.append({"role": "user", "content": user_msg})
        with st.chat_message("assistant"):
            with st.spinner("Thinking…"):
                try:
                    context = st.session_state.paper or st.session_state.all_sections or "No paper content yet."
                    chat_prompt = (
                        f"You are an academic assistant. Answer the user concisely and cite parts of the provided paper when relevant.\n\n"
                        f"User question: {user_msg}\n\n"
                        f"Paper context:\n{context}"
                    )
                    reply = model.generate_content(chat_prompt).text
                except Exception as e:
                    reply = f"Error from Gemini: {e}"
                st.markdown(reply)
        st.session_state.chat_messages.append({"role": "assistant", "content": reply})
